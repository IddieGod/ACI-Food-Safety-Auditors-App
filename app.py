import streamlit as st
import pandas as pd
import base64
import matplotlib.pyplot as plt
import seaborn as sns
from tempfile import NamedTemporaryFile
from docx import Document

# Auditor authentication
def authenticate_auditor():
    auditor_name = st.text_input("Enter your name:", "")

    if auditor_name.lower() in ["callistus kyire", "iddriss nyande", "lovia", "felix"]:
        password = st.text_input("Enter password:", type="password")

        if password == "ACL101":
            st.session_state['auditor_name'] = auditor_name.lower()
            return auditor_name.lower()

    st.error("Invalid auditor name or password.")
    return None

# Home page
def home_page():
    st.title("Airways Catering Limited Audit App")
    st.write("Welcome to the Airways Catering Limited Audit App. Please authenticate to proceed.")

    auditor_name = authenticate_auditor()

    if auditor_name:
        st.success(f"Welcome, {auditor_name.title()}!")
        st.write("App Instructions:")
        st.write("1. Navigate to the Audit Page to conduct audits for your assigned locations.")
        st.write("2. Review the analysis and download the CSV file on the Analysis Page.")
        st.write("3. Add comments and sign out on the Comments and Sign-out Page.")
        st.write("4. Use the navigation sidebar to access different pages.")

# Audit page
def audit_page():
    st.title("Audit Page")
    if 'auditor_name' not in st.session_state:
        st.warning("Please authenticate first on the Home page.")
        return

    auditor_name = st.session_state['auditor_name']
    assigned_food_production_zones = get_assigned_food_production_zones(auditor_name)
    audit_data = []

     # Title Page
    st.header("Auditor Information")
    client_site = st.text_input("Client / Site:", "")
    location = st.text_input("Location:", "")
    address = st.text_input("Position at ACL:", "")
    conducted_on_date = st.date_input("Conducted on:", value=None)

    prepared_by = auditor_name.title()

    st.write("---")

    select_box = st.selectbox("Select:", ["Location", "Food Production Zone"])

    if select_box == "Location":
        for location in ["Exterior", "Interior"]:
            with st.expander(location):
                location_data = conduct_audit_location(location)
                audit_data.extend(location_data)
    elif select_box == "Food Production Zone":
        for zone in assigned_food_production_zones:
            location_data = conduct_audit_food_production_zone(zone)
            if location_data:
                with st.expander(zone):
                    audit_data.extend(location_data)

    st.session_state['audit_data'] = audit_data

def get_assigned_food_production_zones(auditor_name):
    assignments = {
        "callistus kyire": ["Old Lay-up", "Tray Set-up", "Bakery", "Cooked Food Fridge", "Dish Wash-Up Bay"],
        "iddriss nyande": ["Entrance", "Receiving Bay", "Loading Bay", "Old Lay-up Holding Room"],
        "lovia": ["Dry Goods Store", "Hot Kitchen", "Dishing Room", "Butchery", "Pots and Pans Washing Bay"],
        "felix": ["Blast Freezers", "Deep Freezer", "Cold Room"]
    }
    return assignments.get(auditor_name.lower(), [])


def conduct_audit_location(location):
    st.write(f"Conducting audit for {location.title()}")

    questions = get_questions_for_location(location)

    location_data = []

    for i, question in enumerate(questions):
        answer = st.selectbox(question['question'], ["Select Option", "Yes", "No", "N/A"], key=f"{location}_{i}_{question['question']}")
        location_data.append({
            'Location': location,
            'Question': question['question'],
            'Answer': answer
        })

    if st.button(f"Attach photos for {location}", key=f"attach_photos_{location}"):
        uploaded_files = st.file_uploader(f"Upload photos for {location}", accept_multiple_files=True, key=f"upload_photos_{location}")

        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.getvalue()
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
            location_data.append({
                'Location': location,
                'Photo': file_base64
            })

    comments = st.text_area(f"Comments for {location}", key=f"comments_{location}")
    if comments:
        location_data.append({
            'Location': location,
            'Comments': comments
        })

    if not any('Comments' in item for item in location_data):
        location_data.append({
            'Location': location,
            'Comments': None
        })

    if st.button("Submit", key=f"submit_{location}"):
        st.session_state['audit_data'].extend(location_data)
        st.success("Data submitted successfully!")

    # Button to clear the entry
    if st.button("Clear Entry", key=f"clear_{location}"):
        st.session_state['audit_data'] = [entry for entry in st.session_state['audit_data'] if entry['Location'] != location]
        st.info("Entry cleared successfully!")

    return location_data

def conduct_audit_food_production_zone(zone):
    st.write(f"Conducting audit for {zone}")

    questions = get_questions_for_food_production_zone(zone)

    location_data = []

    for i, question in enumerate(questions):
        answer = st.selectbox(question['question'], ["Select Option", "Yes", "No", "N/A"], key=f"{zone}_fpz_{i}_{question['question']}")
        location_data.append({
            'Location': zone,
            'Food Production Zone': question['question'],
            'Answer': answer
        })

    if st.button(f"Attach photos for {zone}", key=f"attach_photos_{zone}"):
        uploaded_files = st.file_uploader(f"Upload photos for {zone}", accept_multiple_files=True, key=f"upload_photos_{zone}")

        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.getvalue()
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
            location_data.append({
                'Location': zone,
                'Photo': file_base64
            })

    comments = st.text_area(f"Comments for {zone}", key=f"comments_{zone}")
    if comments:
        location_data.append({
            'Location': zone,
            'Comments': comments
        })

    if not any('Comments' in item for item in location_data):
        location_data.append({
            'Location': zone,
            'Comments': None
        })

    if st.button("Submit", key=f"submit_{zone}"):
        st.session_state['audit_data'].extend(location_data)
        st.success("Data submitted successfully!")

    # Button to clear the entry
    if st.button("Clear Entry", key=f"clear_{zone}"):
        st.session_state['audit_data'] = [entry for entry in st.session_state['audit_data'] if entry['Location'] != zone]
        st.info("Entry cleared successfully!")

    return location_data

def conduct_audit_location(location):
    st.write(f"Conducting audit for {location.title()}")

    questions = get_questions_for_location(location)

    location_data = []

    for i, question in enumerate(questions):
        answer = st.selectbox(question['question'], ["Yes", "No", "N/A"], key=f"{location}_{i}_{question['question']}")
        location_data.append({
            'Location': location,
            'Question': question['question'],
            'Answer': answer
        })

    if st.button(f"Attach photos for {location}", key=f"attach_photos_{location}"):
        uploaded_files = st.file_uploader(f"Upload photos for {location}", accept_multiple_files=True, key=f"upload_photos_{location}")

        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.getvalue()
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
            location_data.append({
                'Location': location,
                'Photo': file_base64
            })

    comments = st.text_area(f"Comments for {location}", key=f"comments_{location}")
    if comments:
        location_data.append({
            'Location': location,
            'Comments': comments
        })

    if not any('Comments' in item for item in location_data):
        location_data.append({
            'Location': location,
            'Comments': None
        })

    if st.button("Submit", key=f"submit_{location}"):
        st.session_state['audit_data'].extend(location_data)
        st.success("Data submitted successfully!")

    # Button to clear the entry
    if st.button("Clear Entry", key=f"clear_{location}"):
        st.session_state['audit_data'] = [entry for entry in st.session_state['audit_data'] if entry['Location'] != location]
        st.info("Entry cleared successfully!")

    return location_data

def conduct_audit_food_production_zone(zone):
    st.write(f"Conducting audit for {zone}")

    questions = get_questions_for_food_production_zone(zone)

    location_data = []

    for i, question in enumerate(questions):
        answer = st.selectbox(question['question'], ["Yes", "No", "N/A"], key=f"{zone}_fpz_{i}_{question['question']}")
        location_data.append({
            'Location': zone,
            'Food Production Zone': question['question'],
            'Answer': answer
        })

    if st.button(f"Attach photos for {zone}", key=f"attach_photos_{zone}"):
        uploaded_files = st.file_uploader(f"Upload photos for {zone}", accept_multiple_files=True, key=f"upload_photos_{zone}")

        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.getvalue()
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
            location_data.append({
                'Location': zone,
                'Photo': file_base64
            })

    comments = st.text_area(f"Comments for {zone}", key=f"comments_{zone}")
    if comments:
        location_data.append({
            'Location': zone,
            'Comments': comments
        })

    if not any('Comments' in item for item in location_data):
        location_data.append({
            'Location': zone,
            'Comments': None
        })

    if st.button("Submit", key=f"submit_{zone}"):
        st.session_state['audit_data'].extend(location_data)
        st.success("Data submitted successfully!")

    # Button to clear the entry
    if st.button("Clear Entry", key=f"clear_{zone}"):
        st.session_state['audit_data'] = [entry for entry in st.session_state['audit_data'] if entry['Location'] != zone]
        st.info("Entry cleared successfully!")

    return location_data


def get_questions_for_location(location):
    if location == "Exterior":
        questions = [
            {'question': 'Is signage in place, in good repair and clearly visible?'},
            {'question': 'Does the landscaping appear to be clean and well maintained?'},
            {'question': 'Do the sidewalks appear to be clean and in good condition?'},
            {'question': 'Does the building appear to be clean and in good condition? (including windows)'}
        ]

        return questions

    elif location == "Interior":
        questions = [
            {'question': 'Are the restrooms clean and stocked? Free of odor? Seats Secure?'},
            {'question': 'Are the floors, tables and seating clean and in good condition?'},
            {'question': 'Are the trash containers clean and in good condition? Liners tucked? No odors?'},
            {'question': 'Are the front corridors clean and clutter free?'},
            {'question': 'Is all lighting working and in good repair?'},
            {'question': 'Is the area clean and well maintained?'}
        ]

        return questions

    else:
        return get_questions_for_food_production_zone(location)

def get_questions_for_food_production_zone(zone):
    questions = {
        "Entrance": [
            {'question': 'Are floors clean, dry and in good condition?'},
            {'question': 'Are entrance doors and latches working properly?'},
            {'question': 'Is the hand washing station operational?'},
            {'question': 'Are aisles free of obstructions?'},
            {'question': 'Is the temperature within the acceptable range?'},
            {'question': 'Are safety signs and instructions clearly visible?'},
            {'question': 'Are entry and exit points well-marked and maintained?'},
            {'question': 'Is there adequate lighting in the entrance area?'}
        ],
        "Old Lay-up": [
            {'question': 'Is the holding room temperature maintained between 0-5 degrees Celsius?'},
            {'question': 'Are portioned foods properly stored and labeled?'},
            {'question': 'Is the room free of any signs of pest infestation?'},
            {'question': 'Are hygiene protocols followed during food handling?'},
            {'question': 'Is equipment clean and in good condition?'},
            {'question': 'Are temperature logs regularly maintained and reviewed?'},
            {'question': 'Are emergency exits and equipment accessible and functional?'},
            {'question': 'Is there sufficient space for staff to work safely?'}
        ],
        "Dry Goods Store": [
            {'question': 'Are dry goods properly stored in sealed containers?'},
            {'question': 'Is the room well-ventilated to prevent moisture buildup?'},
            {'question': 'Are allergens stored separately and clearly labeled?'},
            {'question': 'Are shelves organized and free of spills?'},
            {'question': 'Are pest control measures in place and effective?'},
            {'question': 'Is there adequate space for maneuvering and storage?'},
            {'question': 'Are spill containment measures in place and functional?'},
            {'question': 'Is there a designated area for receiving and inspecting goods?'}
        ],
        "Hot kitchen": [
            {'question': 'Are cooking utensils and equipment clean and sanitized?'},
            {'question': 'Is food cooked to the required temperature?'},
            {'question': 'Are staff following proper food safety protocols?'},
            {'question': 'Are waste bins emptied regularly and kept covered?'},
            {'question': 'Is the kitchen well-ventilated to remove cooking odors?'},
            {'question': 'Are cooking surfaces and equipment free of grease buildup?'},
            {'question': 'Is there sufficient space for staff movement and workflow?'},
            {'question': 'Are emergency shutdown procedures clearly posted and understood?'}
        ],
        "Dishing Room": [
            {'question': 'Is food properly cooled before dishing?'},
            {'question': 'Are disposable bowls stored in a clean and dry area?'},
            {'question': 'Is there sufficient space to work safely?'},
            {'question': 'Are staff wearing appropriate personal protective equipment?'},
            {'question': 'Are hygiene standards maintained during food portioning?'},
            {'question': 'Are portioning equipment and utensils clean and sanitized?'},
            {'question': 'Is there adequate lighting for accurate food inspection?'},
            {'question': 'Are food handling procedures clearly documented and followed?'}
        ],
        "Butchery": [
            {'question': 'Is meat stored at the correct temperature?'},
            {'question': 'Are cutting boards and knives sanitized between uses?'},
            {'question': 'Is cross-contamination prevented during meat preparation?'},
            {'question': 'Are meat products properly labeled with dates and types?'},
            {'question': 'Are staff trained in safe meat handling practices?'},
            {'question': 'Is there adequate ventilation to remove meat processing odors?'},
            {'question': 'Are meat storage areas organized and free of spills?'},
            {'question': 'Is there a designated area for waste disposal and storage?'}
        ],
        "Blast Freezers": [
            {'question': 'Are blast freezers operating at the correct temperature?'},
            {'question': 'Is food properly packaged before entering the blast freezer?'},
            {'question': 'Are blast freezer doors kept closed when not in use?'},
            {'question': 'Is there adequate space for airflow within the freezer?'},
            {'question': 'Are temperature logs maintained and reviewed regularly?'},
            {'question': 'Is there a backup power source in case of power failure?'},
            {'question': 'Are blast freezer surfaces clean and free of ice buildup?'},
            {'question': 'Are emergency alarms and shutdown procedures in place?'}
        ],
        "Deep Freezer": [
            {'question': 'Is the deep freezer operating at the correct temperature?'},
            {'question': 'Are frozen foods properly stored and organized?'},
            {'question': 'Are freezer shelves free of frost buildup?'},
            {'question': 'Is there a backup power source in case of a power outage?'},
            {'question': 'Are temperature alarms functioning correctly?'},
            {'question': 'Is there a designated area for inventory management?'},
            {'question': 'Are freezer surfaces clean and free of spills?'},
            {'question': 'Are freezer doors and seals well-maintained and functional?'}
        ],
        "Cold Room": [
            {'question': 'Is the cold room temperature within the acceptable range?'},
            {'question': 'Are fruits and vegetables stored separately to prevent cross-contamination?'},
            {'question': 'Are shelves and storage bins clean and free of spills?'},
            {'question': 'Is there adequate lighting in the cold room?'},
            {'question': 'Are temperature logs maintained and reviewed regularly?'},
            {'question': 'Is there a backup cooling system in case of failure?'},
            {'question': 'Are emergency exits and pathways clearly marked and accessible?'},
            {'question': 'Is there sufficient space for staff movement and storage?'}
        ],
        "Tray Set-up": [
            {'question': 'Are food trays clean and sanitized before use?'},
            {'question': 'Is food arranged on trays according to standard procedures?'},
            {'question': 'Are tray assembly areas free of spills and debris?'},
            {'question': 'Are trays inspected for quality before distribution?'},
            {'question': 'Are trays stored in a clean and dry environment?'},
            {'question': 'Are tray assembly areas well-ventilated and lit?'},
            {'question': 'Are tray assembly procedures clearly documented and followed?'},
            {'question': 'Are there designated areas for tray assembly and storage?'}
        ],
        "Bakery": [
            {'question': 'Are baking ingredients stored in airtight containers?'},
            {'question': 'Is baking equipment clean and in good working condition?'},
            {'question': 'Are baked goods cooled properly before storage?'},
            {'question': 'Are bakery products labeled with expiration dates?'},
            {'question': 'Are hygiene standards maintained during baking operations?'},
            {'question': 'Is there sufficient space for equipment maintenance and storage?'},
            {'question': 'Are bakery waste disposal procedures followed properly?'},
            {'question': 'Is there a backup plan for oven and mixer breakdowns?'}
        ],
        "Cooked Food Fridge": [
            {'question': 'Is the fridge temperature maintained between 0-5 degrees Celsius?'},
            {'question': 'Are cooked foods properly covered and labeled?'},
            {'question': 'Is the fridge organized to prevent cross-contamination?'},
            {'question': 'Are temperature logs maintained and reviewed regularly?'},
            {'question': 'Are fridge shelves clean and free of spills?'},
            {'question': 'Is there a backup cooling system in case of failure?'},
            {'question': 'Is there sufficient space for inventory management?'},
            {'question': 'Are emergency shutdown procedures clearly posted and understood?'}
        ],
       "Receiving Bay": [
    {'question': 'Are temperature-sensitive products properly stored upon arrival?'},
    {'question': 'Is there a designated area for receiving and inspecting goods?'},
    {'question': 'Are incoming deliveries properly documented and logged?'},
    {'question': 'Are hygiene standards maintained during product handling?'},
    {'question': 'Is there adequate space for unloading and storage?'},
    {'question': 'Are pest control measures in place and effective?'},
    {'question': 'Is there a backup plan for receiving area breakdowns?'}
],
        "Loading Bay": [
            {'question': 'Is the loading bay area clean and free of spills?'},
            {'question': 'Are outgoing food products properly packaged and labeled?'},
            {'question': 'Is there a designated area for loading food products onto vehicles?'},
            {'question': 'Are loading schedules coordinated to minimize delays?'},
            {'question': 'Are temperature-sensitive products monitored during loading?'},
            {'question': 'Is there adequate space for vehicle maneuvering and parking?'},
            {'question': 'Are safety procedures followed during loading operations?'},
            {'question': 'Is there a backup plan for loading bay breakdowns?'}
        ],
        "Dish Wash-Up Bay": [
            {'question': 'Are dishwashing machines properly maintained and sanitized?'},
            {'question': 'Is there a backup plan in case of dishwasher failure?'},
            {'question': 'Are dishwashing areas kept clean and organized?'},
            {'question': 'Is there adequate ventilation to remove steam and heat?'},
            {'question': 'Are dishwashing detergents and sanitizers used correctly?'},
            {'question': 'Are dishes and utensils properly dried after washing?'},
            {'question': 'Is there sufficient space for dish storage and drying racks?'},
            {'question': 'Are dishwashing schedules followed consistently?'}
        ],
        "Pots and Pans Washing Bay": [
            {'question': 'Are pots and pans properly cleaned and sanitized?'},
            {'question': 'Are cleaning agents used according to safety guidelines?'},
            {'question': 'Is there adequate ventilation in the washing bay?'},
            {'question': 'Are washed pots and pans properly dried before storage?'},
            {'question': 'Is there a backup plan in case of equipment failure?'},
            {'question': 'Are cleaning schedules followed consistently?'},
            {'question': 'Is there sufficient space for equipment maneuvering?'},
            {'question': 'Are pot and pan storage areas clean and organized?'}
        ],
        "Old Lay-up Holding Room": [
            {'question': 'Is the holding room temperature maintained between 0-5 degrees Celsius?'},
            {'question': 'Are portioned foods properly stored and labeled?'},
            {'question': 'Is the room free of any signs of pest infestation?'},
            {'question': 'Are hygiene protocols followed during food handling?'},
            {'question': 'Is equipment clean and in good condition?'},
            {'question': 'Are temperature logs maintained and reviewed regularly?'},
            {'question': 'Are emergency exits and equipment accessible and functional?'},
            {'question': 'Is there sufficient space for staff to work safely?'}
        ],
        # Add more zones and their respective questions here
    }
    return questions.get(zone, [])


def analysis_page():
    # Disable PyplotGlobalUseWarning
    st.set_option('deprecation.showPyplotGlobalUse', False)
    st.title("Analysis Page")
    if 'audit_data' in st.session_state:
        audit_data = st.session_state['audit_data']
        df = pd.DataFrame(audit_data)
        
        st.write(df)
        
        # Add your code for displaying visualizations
        st.subheader("Visualization")
        
        # Add select box for choosing between Location and Food Production Zone
        select_box = st.selectbox("Select:", ["Location", "Food Production Zone"])
        
        # Add clear dataframe button
        if st.button("Clear DataFrame"):
            st.session_state.pop('audit_data', None)
            st.info("DataFrame cleared successfully!")
            return
        
        if select_box == "Location":
            # Example visualization: Bar plot of audit results by location
            st.write("Audit Results by Location:")
            location_counts = df['Location'].value_counts()
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.barplot(x=location_counts.index, y=location_counts.values, ax=ax)
            ax.set_xlabel('Location')
            ax.set_ylabel('Count')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
            # Pie chart of audit results
            st.write("Pie Chart of Audit Results:")
            pie_data = df['Answer'].value_counts()
            fig, ax = plt.subplots(figsize=(8, 8))
            ax.pie(pie_data, labels=pie_data.index, autopct='%1.1f%%', startangle=140)
            ax.axis('equal')
            st.pyplot(fig)
            
            # Histogram of audit data
            st.write("Histogram of Audit Data:")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.histplot(data=df, x='Location', hue='Answer', multiple='stack', ax=ax)
            ax.set_xlabel('Location')
            ax.set_ylabel('Count')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
            # Scatter plot of audit data
            st.write("Scatter Plot of Audit Data:")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.scatterplot(data=df, x='Location', y='Answer', hue='Answer', ax=ax)
            ax.set_xlabel('Location')
            ax.set_ylabel('Answer')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
            # Box plot of audit data
            st.write("Box Plot of Audit Data:")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.boxplot(data=df, x='Location', y='Answer', ax=ax)
            ax.set_xlabel('Location')
            ax.set_ylabel('Answer')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
        elif select_box == "Food Production Zone":
            # Visualizations for Food Production Zone
            # Bar plot of audit results by food production zone
            st.write("Audit Results by Food Production Zone:")
            fpz_counts = df['Food Production Zone'].value_counts()
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.barplot(x=fpz_counts.index, y=fpz_counts.values, ax=ax)
            ax.set_xlabel('Food Production Zone')
            ax.set_ylabel('Count')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
            # Pie chart of audit results by food production zone
            st.write("Pie Chart of Audit Results by Food Production Zone:")
            pie_data = df['Answer'].value_counts()
            fig, ax = plt.subplots(figsize=(8, 8))
            ax.pie(pie_data, labels=pie_data.index, autopct='%1.1f%%', startangle=140)
            ax.axis('equal')
            st.pyplot(fig)
            
            # Histogram of audit results by food production zone
            st.write("Histogram of Audit Results by Food Production Zone:")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.histplot(data=df, x='Food Production Zone', hue='Answer', multiple='stack', ax=ax)
            ax.set_xlabel('Food Production Zone')
            ax.set_ylabel('Count')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
            # Box plot of audit results by food production zone
            st.write("Box Plot of Audit Results by Food Production Zone:")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.boxplot(data=df, x='Food Production Zone', y='Answer', ax=ax)
            ax.set_xlabel('Food Production Zone')
            ax.set_ylabel('Answer')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45)
            st.pyplot(fig)
            
        if st.button("Download CSV"):
            csv = df.to_csv(index=False)
            b64 = base64.b64encode(csv.encode()).decode()
            href = f'<a href="data:file/csv;base64,{b64}" download="audit_data.csv">Download CSV File</a>'
            st.markdown(href, unsafe_allow_html=True)



def comments_sign_out_page():
    st.title("Comments and Sign-out Page")
    if 'audit_data' in st.session_state:
        audit_data = st.session_state['audit_data']
        comments = [row.get('Comments', '') for row in audit_data if 'Comments' in row]

        if comments:
            st.write("Comments:")
            for comment in comments:
                st.write(comment)
            # Create a Word document to save comments
            document = Document()
            document.add_heading('Audit Comments', level=1)
            for comment in comments:
                document.add_paragraph(comment)
            st.markdown(get_word_download_link(document), unsafe_allow_html=True)
        else:
            st.write("No comments found.")
        
        # Text input for adding new comments
        new_comment = st.text_area("Add a new comment:")
        
        signature = st.text_input("Enter your signature:")
        if st.button("Sign Out"):
            del st.session_state['audit_data']
            st.success("Signed out successfully!")


def get_word_download_link(document):
    """
    Generate a download link for the Word document.
    """
    tmp_download_link = None
    with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        document.save(tmp_file.name)
        tmp_download_link = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(open(tmp_file.name, "rb").read()).decode()}">Download Comments as Word Document</a>'
    return tmp_download_link


# Main app
def main():
    # Increase font size
    st.markdown(
        """
        <style>
        .css-1aumxhk {
            font-size: 18px !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    pages = {
        "Home": home_page,
        "Audit": audit_page,
        "Analysis": analysis_page,
        "Comments and Sign-out": comments_sign_out_page
    }

    st.sidebar.title("ACL Food Safety Auditor")
    selection = st.sidebar.radio("Go to", list(pages.keys()))

    pages[selection]()

if __name__ == "__main__":
    main()
